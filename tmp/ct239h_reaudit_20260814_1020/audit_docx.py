import json
import re
import sys
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
}


def qn(ns, tag):
    return f"{{{NS[ns]}}}{tag}"


path = Path(sys.argv[1])
doc = Document(path)
paras = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        paras.append({"i": i, "style": p.style.name, "text": text})

tables = []
for ti, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip().replace("\n", " | ") for cell in row.cells])
    tables.append({"i": ti, "rows": rows})

with zipfile.ZipFile(path) as zf:
    styles = ET.fromstring(zf.read("word/styles.xml"))
    theme = ET.fromstring(zf.read("word/theme/theme1.xml"))
    document_xml = ET.fromstring(zf.read("word/document.xml"))

theme_fonts = {}
font_scheme = theme.find(".//a:fontScheme", NS)
if font_scheme is not None:
    for kind in ("majorFont", "minorFont"):
        node = font_scheme.find(f"a:{kind}", NS)
        if node is not None:
            theme_fonts[kind] = {
                "latin": (node.find("a:latin", NS) or {}).get("typeface"),
                "ea": (node.find("a:ea", NS) or {}).get("typeface"),
                "cs": (node.find("a:cs", NS) or {}).get("typeface"),
            }

style_info = {}
for st in styles.findall("w:style", NS):
    name_node = st.find("w:name", NS)
    name = name_node.get(qn("w", "val")) if name_node is not None else ""
    if name.lower() not in {"normal", "heading 1", "heading 2", "heading 3", "heading 4"}:
        continue
    rpr = st.find("w:rPr", NS)
    fonts = {}
    size = None
    if rpr is not None:
        rf = rpr.find("w:rFonts", NS)
        if rf is not None:
            fonts = {k.split("}")[-1]: v for k, v in rf.attrib.items()}
        sz = rpr.find("w:sz", NS)
        if sz is not None:
            size = sz.get(qn("w", "val"))
    style_info[name] = {"styleId": st.get(qn("w", "styleId")), "fonts": fonts, "half_points": size}

fields = []
for instr in document_xml.findall(".//w:instrText", NS):
    if instr.text:
        fields.append(instr.text.strip())

out = {
    "path": str(path),
    "paragraph_count": len(doc.paragraphs),
    "nonempty_count": len(paras),
    "table_count": len(doc.tables),
    "section_count": len(doc.sections),
    "inline_shapes": len(doc.inline_shapes),
    "styles_used": Counter(p["style"] for p in paras),
    "paragraphs": paras,
    "tables": tables,
    "style_info": style_info,
    "theme_fonts": theme_fonts,
    "fields": fields,
    "tracked_insertions": len(document_xml.findall(".//w:ins", NS)),
    "tracked_deletions": len(document_xml.findall(".//w:del", NS)),
}
print(json.dumps(out, ensure_ascii=False, indent=2, default=dict))
