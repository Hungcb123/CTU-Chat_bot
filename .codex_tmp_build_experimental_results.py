from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"D:\Project\Chatbot\Báo_cáo\Ket_qua_chay_thuc_nghiem_CT239H.docx")
FONT = "Times New Roman"
TABLE_WIDTH = 8951
TABLE_INDENT = 120


def set_run_font(run, size=13, bold=None, italic=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, (cell, width) in enumerate(zip(row.cells, widths)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def format_cell(cell, header=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        for run in p.runs:
            set_run_font(run, size=size, bold=header)
    if header:
        set_cell_shading(cell, "D9E2F3")


def add_table(doc, caption, headers, rows, widths, alignments=None, font_size=11):
    cp = doc.add_paragraph(style="Caption")
    cp.paragraph_format.keep_with_next = True
    cp.add_run(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    repeat_header(hdr)
    prevent_row_split(hdr)
    for idx, value in enumerate(headers):
        hdr.cells[idx].text = value
        format_cell(hdr.cells[idx], header=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=font_size)
    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, value in enumerate(row_values):
            row.cells[idx].text = str(value)
            alignment = (alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
            format_cell(row.cells[idx], align=alignment, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    p.add_run(text)
    return p


def add_bullet(doc, label, detail):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label)
    set_run_font(r, bold=True)
    r = p.add_run(detail)
    set_run_font(r)
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1.27)

    settings = {
        "Heading 1": (14, True, False, 16, 8),
        "Heading 2": (13, True, False, 13, 5),
        "Heading 3": (13, True, True, 10, 4),
    }
    for name, (size, bold, italic, before, after) in settings.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)

    caption = styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption.font.size = Pt(13)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.line_spacing = 1.0

    lb = styles["List Bullet"]
    lb.font.name = FONT
    lb._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    lb._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    lb.font.size = Pt(13)
    lb.paragraph_format.left_indent = Cm(0.75)
    lb.paragraph_format.first_line_indent = Cm(-0.35)
    lb.paragraph_format.space_after = Pt(4)
    lb.paragraph_format.line_spacing = 1.35


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(14)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("EXPERIMENTAL RESULTS"), size=16, bold=True)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    lead.paragraph_format.space_after = Pt(10)
    lead.paragraph_format.line_spacing = 1.5
    lead.paragraph_format.first_line_indent = Cm(1.27)
    set_run_font(lead.add_run(
        "This document consolidates the two empirical evaluations conducted for the CTU student-finance chatbot: "
        "a retrieval-only reranker comparison and an authenticated end-to-end chatbot evaluation. The results are "
        "reported separately because retrieval quality and final-answer quality measure different parts of the system."
    ))

    add_heading(doc, "1. Experimental objectives", 1)
    add_body(doc,
        "The experiments were designed to answer two practical questions. First, which reranker provides the best "
        "retrieval quality under the available hardware constraint? Second, how accurately does the complete chatbot "
        "answer representative Vietnamese student-finance questions when it is exercised through the same HTTP boundary "
        "used by the browser?"
    )
    add_bullet(doc, "Retrieval experiment. ", "Measure whether the required source and fact appear within the six parent passages supplied to generation.")
    add_bullet(doc, "End-to-end experiment. ", "Measure whether the final chatbot answer captures the essential facts in the expected answer.")

    add_heading(doc, "2. Retrieval and reranker experiment", 1)
    add_heading(doc, "2.1. Experimental setup", 2)
    add_body(doc,
        "A controlled retrieval-only comparison was conducted on ten difficult questions drawn from tuition exemption, "
        "social support, student-loan, and scholarship scenarios. Gemini was excluded from answer generation and judging "
        "during this comparison. Both configurations used the same Vietnamese bi-encoder, Qdrant index, query set, and "
        "candidate passages. The first stage retrieved up to 15 candidate parent passages, and the reranker retained the "
        "six highest-ranked passages. Each query-passage pair was limited to 512 tokens."
    )

    add_table(
        doc,
        "Table 1. Retrieval-only experimental configuration",
        ["Item", "Configuration"],
        [
            ["Evaluation set", "10 difficult Vietnamese student-finance questions"],
            ["Shared first stage", "Vietnamese Bi-Encoder and Qdrant"],
            ["Candidate depth", "Top 15 parent passages"],
            ["Final context depth", "Top 6 reranked parent passages"],
            ["Pair length limit", "512 tokens"],
            ["Hardware", "NVIDIA GTX 1650 GPU with 4 GB VRAM"],
            ["Temporal rule", "Prefer newer content only when reranker scores differ by no more than 0.05"],
        ],
        [2400, 6551],
        font_size=11,
    )

    add_heading(doc, "2.2. Metrics", 2)
    add_body(doc,
        "Hit@6 was counted only when the final six parent passages contained both the expected source and the specific "
        "fact required to answer the question. Mean Reciprocal Rank (MRR) measured how highly the first relevant parent "
        "passage was ranked. These retrieval-only metrics were treated as the primary evidence for model selection."
    )

    add_heading(doc, "2.3. Results and interpretation", 2)
    add_table(
        doc,
        "Table 2. Comparison of BGE and GTE rerankers",
        ["Criterion", "BGE Reranker v2-m3", "GTE Multilingual Reranker Base"],
        [
            ["Shared evaluation set", "10 questions", "10 questions"],
            ["Hit@6", "70%", "60%"],
            ["Mean Reciprocal Rank", "0.65", "0.50"],
            ["End-to-end accuracy", "95%", "94%"],
            ["Main observation", "Better ranking of fact-bearing parent passages", "Faster, but lower retrieval quality"],
        ],
        [3000, 2825, 3126],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=10.5,
    )
    add_body(doc,
        "BGE Reranker v2-m3 produced the stronger retrieval-only result, improving Hit@6 by ten percentage points and "
        "MRR by 0.15. The one-percentage-point end-to-end accuracy difference was considered supporting rather than primary "
        "evidence because final answers can also be affected by Gemini generation and LLM-as-a-judge variability. BGE was "
        "therefore selected as the most appropriate trade-off for this dataset and hardware configuration; the result does "
        "not establish universal superiority over GTE or other rerankers."
    )
    add_table(
        doc,
        "Table 3. Retrieval and reranking responsibilities",
        ["Stage", "Model or method", "Input", "Output", "Primary goal"],
        [
            ["1", "Vietnamese Bi-Encoder + Qdrant", "Question and precomputed child vectors", "Top 15 candidate parents", "Fast recall"],
            ["2", "BGE Reranker v2-m3", "Question-parent pairs", "Top 6 ordered parents", "Fine-grained precision"],
            ["3", "Soft temporal tie-break", "Near-equal reranker scores", "Stable recency-aware order", "Prefer current policy without overriding relevance"],
        ],
        [700, 2100, 2050, 1800, 2301],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=9.5,
    )

    add_heading(doc, "3. End-to-end chatbot evaluation", 1)
    add_heading(doc, "3.1. Dataset and execution procedure", 2)
    add_body(doc,
        "The evaluation dataset contains 100 Vietnamese question-answer cases distributed equally across tuition exemption "
        "and learning-cost support, tuition, student loans, and scholarships. Each case contains a question, an expected "
        "answer, and one or more expected source documents. Expected sources are stored for traceability but are not included "
        "in the automated score."
    )
    add_body(doc,
        "The evaluator authenticates once through POST /auth/login, retains the HTTP-only cookie, and submits every question "
        "to POST /chat. By default, each case starts without a shared session identifier so that previous cases do not affect "
        "the current answer. This procedure includes authentication, request validation, intent routing, retrieval, reranking, "
        "Gemini generation, tool calling when required, and response serialization."
    )
    add_table(
        doc,
        "Table 4. End-to-end evaluation configuration",
        ["Item", "Recorded value"],
        [
            ["Run started", "31 July 2026, 18:59:40 (UTC+07:00)"],
            ["Dataset size", "100 questions; 25 questions per domain"],
            ["Application endpoint", "Authenticated POST /chat"],
            ["Judge model", "Gemini 3.1 Flash-Lite"],
            ["Pass threshold", "0.55 on a 0.0-1.0 score"],
            ["Session policy", "Independent session by default"],
            ["HTTP/API errors", "0"],
        ],
        [2400, 6551],
        font_size=11,
    )

    add_heading(doc, "3.2. Scoring method", 2)
    add_body(doc,
        "Gemini 3.1 Flash-Lite acts as an LLM-as-a-judge. For each case, it compares the actual response with the expected "
        "answer, assigns a score from 0.0 to 1.0, returns a pass/fail decision, and provides a short explanation. A case is "
        "accepted only when the judge marks it as passed, the score is at least 0.55, and the response captures the essential "
        "expected facts. Missing answers and API failures receive a score of 0.0. Accuracy is the proportion of passed cases; "
        "average score is the arithmetic mean of all judge scores."
    )

    add_heading(doc, "3.3. Results", 2)
    add_table(
        doc,
        "Table 5. Chatbot accuracy evaluation results",
        ["Domain", "Passed", "Total", "Accuracy", "Average score"],
        [
            ["Tuition exemption and learning-cost support", "25", "25", "100%", "99.20%"],
            ["Tuition", "23", "25", "92%", "94.00%"],
            ["Student loans", "24", "25", "96%", "93.00%"],
            ["Scholarships", "23", "25", "92%", "92.40%"],
            ["Overall", "95", "100", "95%", "94.65%"],
        ],
        [3700, 950, 950, 1500, 1851],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=10.5,
    )
    add_body(doc,
        "The system passed 95 of 100 questions and obtained an average judge score of 94.65%. Tuition exemption and learning-cost "
        "support achieved the strongest result, with all 25 cases passing. Tuition and scholarship questions each recorded two "
        "failed cases, while student loans recorded one failed case. The domain averages remained above 92%, indicating broadly "
        "consistent performance across the balanced test set."
    )

    add_heading(doc, "4. Limitations of the experiments", 1)
    add_body(doc,
        "The retrieval-only comparison uses ten difficult questions and therefore supports a project-specific engineering decision "
        "rather than a general ranking benchmark. The end-to-end evaluation uses an LLM judge, so individual scores may vary with "
        "model behavior and provider updates. The current scorer evaluates semantic agreement with the expected answer but does not "
        "automatically verify whether the response cites or faithfully uses the expected source. Manual review of failed, borderline, "
        "and high-impact financial-policy cases remains necessary."
    )

    add_heading(doc, "5. Experimental conclusion", 1)
    add_body(doc,
        "Within the available GTX 1650 hardware and the current Vietnamese student-finance corpus, BGE Reranker v2-m3 provided the "
        "stronger retrieval quality and was retained for the deployed pipeline. With this configuration, the complete chatbot achieved "
        "95% LLM-judged accuracy and a 94.65% average score on the 100-question balanced dataset. These results support the practical "
        "feasibility of the system while preserving the need for source-grounded review and repeated evaluation after corpus or model changes."
    )

    add_heading(doc, "6. Reproducibility records", 1)
    add_table(
        doc,
        "Table 6. Experimental evidence files",
        ["Artifact", "Project-relative path"],
        [
            ["Evaluation dataset", "data/dataset.md"],
            ["Evaluation script", "scripts/evaluate_chat_dataset.py"],
            ["Human-readable result report", "logs/dataset_evaluation/chat_eval_20260731_185940.md"],
            ["Machine-readable evidence", "logs/dataset_evaluation/chat_eval_20260731_185940.jsonl"],
        ],
        [2900, 6051],
        font_size=10.5,
    )

    props = doc.core_properties
    props.title = "Experimental Results - CT239H Chatbot"
    props.subject = "Retrieval, reranking, and end-to-end chatbot evaluation"
    props.author = "CT239H Project Team"
    props.keywords = "CT239H, chatbot, RAG, reranker, evaluation"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
